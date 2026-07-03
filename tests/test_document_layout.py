from __future__ import annotations

import json
from datetime import date
from pathlib import Path

import pytest

from app.config import Settings
from app.models import Case, User
from app.services.document_qa import run_document_qa
from app.services.document_templates.renderer import _render_statement_docx, create_case_documents
from app.services.document_templates.statement_templates import StatementContext, build_attachments, build_header_lines, build_statement_paragraphs
from app.services.document_templates.styles import A4_HEIGHT, A4_WIDTH, MARGIN_LEFT, StyleProfile
from app.services.document_visual_qa import run_visual_qa
from app.services.legal_data import legal_deadline_from_received, normalize_order_data, validate_amounts
from app.services.pdf_tools import pdf_page_count, pdf_text
from docx import Document

LONG_POST_BANK_ADDRESS = "107061, г. Москва, Преображенская пл., д. 8; 101000, г. Москва, по улице Мясницкая, д. 35; для корреспонденции: 443001, г. Самара ул. Галактионовская д. 157"


BELSKY_DATA = {
    "court_name": "судебный участок №5 города Ессентуки Ставропольского края",
    "court_address": "357600, Ставропольский край, город Ессентуки, улица Шмидта, дом № 72",
    "debtor_name_raw": "Бельскому Владимиру Геннадьевичу",
    "debtor_full_name": "Бельский Владимир Геннадьевич",
    "debtor_address": "г. Ессентуки, ул. Володарского, д. 14, кв. 9",
    "creditor_name": "АО «Почта Банк»",
    "creditor_address": "107061, г. Москва, Преображенская пл., д. 8",
    "case_number": "2-146-09-434/2021",
    "uid": "26MS0031-01-2021-000169-72",
    "order_date": "18.01.2021",
    "debt_contract": "договор № 43006327 от 27 апреля 2019 года",
    "debt_period": "с 27.03.2020 по 28.11.2020",
    "debt_amount": "78 472 руб. 87 коп.",
    "state_duty": "1 277 руб. 00 коп.",
    "total_amount": "79 749 руб. 87 коп.",
}


def _settings(**kwargs):
    base = dict(
        telegram_bot_token="",
        max_bot_token="",
            max_api_base_url="https://platform-api2.max.ru",
            max_use_webhook=False,
            max_webhook_url=None,
            max_webhook_secret=None,
            max_webhook_host="0.0.0.0",
            max_webhook_port=8081,
            max_longpoll_timeout_seconds=30,
            max_download_dir="storage/max",
            max_upload_retry_attempts=5,
            max_upload_retry_base_seconds=1,
            max_admin_ids=set(),
        max_debug_raw_updates=False,
        run_telegram=True,
        run_max=False,
        admin_ids=set(),
        manager_ids=set(),
        database_url="sqlite+aiosqlite:///:memory:",
        drop_pending_updates=True,
        openai_api_key=None,
        openai_base_url="https://api.openai.com/v1",
        vision_model="gpt-5.4-mini",
        text_model="gpt-5.4-mini",
        ai_review_model="gpt-4.1",
        ai_review_fallback_model="gpt-4.1-mini",
        llm_timeout_seconds=90,
        max_ai_review_regenerations=1,
        document_ai_review_mode="shadow",
        admin_debug_to_telegram=False,
        document_price_rub=990,
        document_preview_mode="pdf",
        enable_pdf_preview=True,
        require_pdf_preview_for_payment=False,
        allow_dev_docx_preview=True,
        document_template_version="test",
        show_user_confirmation_step=False,
        yoomoney_receiver=None,
        yoomoney_success_url=None,
        yoomoney_notification_secret=None,
        yookassa_enabled=False,
        yookassa_shop_id=None,
        yookassa_secret_key=None,
        yookassa_return_url=None,
        yookassa_webhook_path="/payments/yookassa",
        yookassa_test_mode=False,
        yookassa_receipt_enabled=True,
        yookassa_vat_code=1,
        yookassa_payment_subject="service",
        yookassa_payment_mode="full_payment",
        yookassa_receipt_description="Подготовка заявления об отмене судебного приказа",
        yookassa_test_customer_email="test@example.com",
        yookassa_tax_system_code=None,
        payment_public_base_url=None,
        payment_web_host="0.0.0.0",
        payment_web_port=8080,
        openai_input_price_per_1m=0.75,
        openai_cached_input_price_per_1m=0.075,
        openai_output_price_per_1m=4.50,
        openai_model_pricing_json="",
        amocrm_base_url="https://example.amocrm.ru",
        amocrm_access_token="token",
        amocrm_enabled=False,
        amocrm_pipeline_name="Судебный приказ",
        amocrm_auto_create_pipeline=False,
        amocrm_auto_create_statuses=True,
        amocrm_attach_files=True,
        amocrm_file_upload_enabled=True,
        amocrm_file_upload_timeout_seconds=30,
        amocrm_debug=False,
        amocrm_rps_limit=5,
        amocrm_pipeline_id=None,
        crm_sync_background=True,
        crm_sync_timeout_seconds=5,
        crm_sync_max_attempts=3,
        crm_sync_retry_base_seconds=2,
        crm_sync_debug=False,
        amount_retry_on_mismatch=True,
        auto_recover_amount_mismatch=True,
        auto_recover_amount_min_confidence=0.75,
        company_name="test",
        manager_contact_text="test",
    )
    base.update(kwargs)
    return Settings(**base)


def _make_case(data: dict | None = None, *, envelope: bool = False) -> tuple[Case, User]:
    received = date(2026, 6, 19)
    user = User(id=1, platform="telegram", platform_user_id="test", telegram_id=1)
    case = Case(
        id=42,
        user_id=1,
        platform="telegram",
        status="processing",
        received_date=received,
        deadline_date=legal_deadline_from_received(received),
        extracted_json=json.dumps(normalize_order_data(data or BELSKY_DATA), ensure_ascii=False),
        envelope_photo_path="storage/photos/env.jpg" if envelope else None,
    )
    return case, user


def _generate(tmp_path, monkeypatch, **case_kwargs):
    monkeypatch.setattr("app.services.document_templates.renderer.DOCUMENT_DIR", tmp_path / "documents")
    case, user = _make_case(**case_kwargs)
    return create_case_documents(case, user, _settings(), restore_reason="\u041f\u0440\u0438\u0447\u0438\u043d\u0430 \u043f\u0440\u043e\u043f\u0443\u0441\u043a\u0430 \u0441\u0440\u043e\u043a\u0430: \u0442\u0435\u0441\u0442\u043e\u0432\u0430\u044f \u0433\u0435\u043d\u0435\u0440\u0430\u0446\u0438\u044f.")


def test_amounts_belsky_are_correct():
    check = validate_amounts(normalize_order_data(BELSKY_DATA))
    assert check.ok
    assert str(check.computed_total) == "79749.87"


def test_amount_mismatch_fails_qa():
    bad = dict(BELSKY_DATA)
    bad["debt_amount"] = "78 742 руб. 00 коп."
    check = validate_amounts(normalize_order_data(bad))
    assert not check.ok
    assert "amount_mismatch" in check.errors


def test_total_amount_computed_from_debt_and_state_duty():
    data = normalize_order_data({"debt_amount": "78 472 руб. 87 коп.", "state_duty": "1 277 руб. 00 коп."})
    assert data["total_amount"] == "79 749 руб. 87 коп."


def test_attachments_with_envelope():
    ctx = StatementContext(
        data=normalize_order_data(BELSKY_DATA),
        received_date=date(2026, 6, 19),
        deadline_date=date(2026, 6, 29),
        document_date=date(2026, 6, 24),
        has_envelope=True,
    )
    items = build_attachments(ctx)
    assert any("конверта" in item for item in items)
    assert any("настоящих возражений" in item for item in items)


def test_attachments_with_manual_date():
    ctx = StatementContext(
        data=normalize_order_data(BELSKY_DATA),
        received_date=date(2026, 6, 19),
        deadline_date=date(2026, 6, 29),
        document_date=date(2026, 6, 24),
        manual_date_only=True,
    )
    items = build_attachments(ctx)
    assert any("при наличии" in item for item in items)
    assert all("настоящего заявления" not in item for item in items)


def test_statement_uses_correct_title():
    ctx = StatementContext(
        data=normalize_order_data(BELSKY_DATA),
        received_date=date(2026, 6, 19),
        deadline_date=date(2026, 6, 29),
        document_date=date(2026, 6, 24),
    )
    paragraphs = build_statement_paragraphs(ctx)
    assert "Считаю требования взыскателя спорными" not in " ".join(paragraphs)
    assert "78 472 руб. 87 коп." in paragraphs[0]
    assert "1 277 руб. 00 коп." in paragraphs[0]
    assert "всего" not in paragraphs[0].lower()


def test_statement_no_old_dispute_text():
    ctx = StatementContext(
        data=normalize_order_data(BELSKY_DATA),
        received_date=date(2026, 6, 19),
        deadline_date=date(2026, 6, 29),
        document_date=date(2026, 6, 24),
    )
    text = " ".join(build_statement_paragraphs(ctx))
    for forbidden in (
        "искового производства",
        "проценты",
        "неустойк",
        "комисс",
        "Считаю требования",
    ):
        assert forbidden not in text


def test_statement_belsky_fits_one_page(tmp_path, monkeypatch):
    artifacts = _generate(tmp_path, monkeypatch)
    if artifacts.full_pdf_path:
        assert pdf_page_count(artifacts.full_pdf_path) == 1


def test_statement_has_a4_margins(tmp_path, monkeypatch):
    artifacts = _generate(tmp_path, monkeypatch)
    doc = Document(str(artifacts.full_docx_path))
    section = doc.sections[0]
    assert abs(section.page_width - A4_WIDTH) <= MARGIN_LEFT * 0.02 + MARGIN_LEFT * 0.001
    assert abs(section.page_height - A4_HEIGHT) <= MARGIN_LEFT * 0.02 + MARGIN_LEFT * 0.001


def test_statement_has_no_placeholder_fields(tmp_path, monkeypatch):
    from app.services.legal_data import docx_text

    artifacts = _generate(tmp_path, monkeypatch)
    text = docx_text(str(artifacts.full_docx_path))
    for token in ("________________", "{{", "None", "уточнить", "20__"):
        assert token not in text


def test_statement_signature_is_filled(tmp_path, monkeypatch):
    from app.services.legal_data import docx_text

    artifacts = _generate(tmp_path, monkeypatch)
    text = pdf_text(artifacts.full_pdf_path) if artifacts.full_pdf_path else ""
    if "Бельский В.Г." not in text:
        text = docx_text(str(artifacts.full_docx_path))
    assert "/Бельский В.Г./" in text or "Бельский В.Г." in text


def test_statement_no_weird_justified_spaces(tmp_path, monkeypatch):
    artifacts = _generate(tmp_path, monkeypatch)
    if artifacts.full_pdf_path:
        visual = run_visual_qa(
            full_docx=artifacts.full_docx_path,
            full_pdf=artifacts.full_pdf_path,
            preview_pdf=artifacts.preview_pdf_path,
            data=normalize_order_data(BELSKY_DATA),
            restore_term=False,
            amount_check=validate_amounts(normalize_order_data(BELSKY_DATA)),
        )
        assert "weird_justified_spaces" not in visual.errors


def test_proshu_not_orphaned_at_page_bottom(tmp_path, monkeypatch):
    from app.services.legal_data import docx_text

    artifacts = _generate(tmp_path, monkeypatch)
    if artifacts.full_pdf_path and pdf_page_count(artifacts.full_pdf_path) == 1:
        text = pdf_text(artifacts.full_pdf_path)
        if "ПРОШУ:" not in text:
            text = docx_text(str(artifacts.full_docx_path))
        assert "ПРОШУ:" in text
        assert ("1. \u041e\u0442\u043c\u0435\u043d\u0438\u0442\u044c" in text) or ("2. \u041e\u0442\u043c\u0435\u043d\u0438\u0442\u044c" in text)


def test_document_qa_amount_mismatch(tmp_path):
    bad = normalize_order_data({**BELSKY_DATA, "debt_amount": "78 742 руб. 00 коп."})
    check = validate_amounts(bad)
    qa = run_document_qa(
        data=bad,
        received_date=date(2026, 6, 19),
        deadline_date=date(2026, 6, 29),
        full_docx=None,
        full_pdf=None,
        preview_pdf=None,
        instruction_docx=None,
        require_preview_pdf=False,
        amount_check=check,
    )
    assert not qa.ok
    assert "amount_mismatch" in qa.bad_tokens



def test_creditor_header_uses_first_normalized_address():
    data = normalize_order_data({**BELSKY_DATA, "creditor_address": LONG_POST_BANK_ADDRESS})
    assert "по улице" not in data["creditor_address"].lower()
    assert "ул. Мясницкая" in data["creditor_address"]
    ctx = StatementContext(
        data=data,
        received_date=date(2026, 6, 19),
        deadline_date=date(2026, 7, 10),
        document_date=date(2026, 6, 30),
    )

    lines = build_header_lines(ctx)

    assert "107061, г. Москва, Преображенская пл., д. 8" in lines
    assert all("Мясницкая" not in line for line in lines)
    assert all("Галактионовская" not in line for line in lines)


def test_in_time_long_creditor_address_compacts_to_one_page(tmp_path, monkeypatch):
    monkeypatch.setattr("app.services.document_templates.renderer.DOCUMENT_DIR", tmp_path / "documents")
    case, user = _make_case({**BELSKY_DATA, "creditor_address": LONG_POST_BANK_ADDRESS})
    case.deadline_date = date(2026, 7, 10)

    artifacts = create_case_documents(case, user, _settings(), restore_reason=None)

    assert artifacts.full_pdf_path is not None
    assert pdf_page_count(artifacts.full_pdf_path) == 1
    assert artifacts.visual_qa is not None
    assert "signature_orphaned_on_page2" not in artifacts.visual_qa.errors


def test_visual_qa_rejects_signature_only_page2(tmp_path):
    fitz = pytest.importorskip("fitz")
    pdf = tmp_path / "orphan_signature.pdf"
    doc = fitz.open()
    page1 = doc.new_page()
    page1.insert_text((72, 72), "ВОЗРАЖЕНИЯ\nПРОШУ:\n1. Отменить судебный приказ")
    page2 = doc.new_page()
    page2.insert_text((72, 72), "«30» июня 2026 г. _____________ /Бельский В.Г./")
    doc.save(pdf)
    doc.close()

    visual = run_visual_qa(
        full_docx=None,
        full_pdf=pdf,
        preview_pdf=None,
        data=normalize_order_data(BELSKY_DATA),
        restore_term=False,
        amount_check=validate_amounts(normalize_order_data(BELSKY_DATA)),
    )

    assert "signature_orphaned_on_page2" in visual.errors


def test_signature_has_one_line_gap_after_attachments(tmp_path):
    profile = StyleProfile.normal()
    ctx = StatementContext(
        data=normalize_order_data(BELSKY_DATA),
        received_date=date(2026, 6, 19),
        deadline_date=date(2026, 7, 10),
        document_date=date(2026, 6, 30),
    )
    docx_path = tmp_path / "signature_gap.docx"

    _render_statement_docx(docx_path, ctx, profile)

    doc = Document(str(docx_path))
    paragraphs = doc.paragraphs
    signature_index = next(index for index, p in enumerate(paragraphs) if "_____________" in p.text)
    signature = paragraphs[signature_index]
    spacer = paragraphs[signature_index - 1]
    previous_text = next(p.text for p in reversed(paragraphs[: signature_index - 1]) if p.text.strip())

    assert previous_text.startswith("3. ")
    assert spacer.text == " "
    assert spacer.paragraph_format.space_before.pt == pytest.approx(0, abs=0.1)
    assert spacer.paragraph_format.space_after.pt == pytest.approx(0, abs=0.1)
    assert spacer.paragraph_format.line_spacing == 1.0
    assert spacer.runs[0].font.size.pt == pytest.approx(profile.body_font_size, abs=0.1)
    assert signature.paragraph_format.space_before.pt == pytest.approx(0, abs=0.1)
    assert signature.paragraph_format.space_after.pt == pytest.approx(0, abs=0.1)

def test_header_uses_clean_registration_address_and_protects_house_number():
    raw_address = (
        "\u0443\u0440\u043e\u0436\u0435\u043d\u0435\u0446 \u0433. \u0410\u0447\u0438\u043d\u0441\u043a \u041a\u0440\u0430\u0441\u043d\u043e\u044f\u0440\u0441\u043a\u043e\u0433\u043e \u043a\u0440\u0430\u044f, "
        "\u0437\u0430\u0440\u0435\u0433\u0438\u0441\u0442\u0440\u0438\u0440\u043e\u0432\u0430\u043d\u043d\u043e\u043c\u0443 \u0432 \u0433\u043e\u0440\u043e\u0434\u0435 \u0415\u0441\u0441\u0435\u043d\u0442\u0443\u043a\u0438, "
        "\u0443\u043b. \u0412\u043e\u043b\u043e\u0434\u0430\u0440\u0441\u043a\u043e\u0433\u043e \u0434. 14, \u043a\u0432. 9"
    )
    data = normalize_order_data({**BELSKY_DATA, "debtor_address": raw_address})
    ctx = StatementContext(
        data=data,
        received_date=date(2026, 6, 19),
        deadline_date=date(2026, 7, 10),
        document_date=date(2026, 6, 30),
    )

    header = "\n".join(build_header_lines(ctx))

    assert "\u0430\u0434\u0440\u0435\u0441: \u0433. \u0415\u0441\u0441\u0435\u043d\u0442\u0443\u043a\u0438, \u0443\u043b. \u0412\u043e\u043b\u043e\u0434\u0430\u0440\u0441\u043a\u043e\u0433\u043e, \u0434. 14, \u043a\u0432. 9" in header
    assert "\u0410\u0447\u0438\u043d\u0441\u043a" not in header
    assert "\u0437\u0430\u0440\u0435\u0433\u0438\u0441\u0442\u0440" not in header.lower()
    assert "\u0443\u0440\u043e\u0436\u0435\u043d" not in header.lower()
    assert "\u043f\u0430\u0441\u043f\u043e\u0440\u0442" not in header.lower()
    assert "\u0443\u043b. \u0428\u043c\u0438\u0434\u0442\u0430, \u0434.\xa072" in header
