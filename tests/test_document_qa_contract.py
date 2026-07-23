from __future__ import annotations

from datetime import date
from pathlib import Path

import pytest
from docx import Document

from app.services.document_qa import output_format_violations, run_document_qa
from app.services.document_render_contract import (
    INTENTIONALLY_NOT_RENDERED_FIELDS,
    RENDER_CONTRACT,
    canonicalize_address_for_qa,
    canonicalize_entity_for_qa,
    select_creditor_address_for_render,
)
from app.services.document_templates.renderer import (
    _build_instruction_doc,
    _render_statement_docx,
)
from app.services.document_templates.statement_templates import (
    StatementContext,
    build_header_lines,
    build_statement_paragraphs,
)
from app.services.document_templates.styles import StyleProfile
from app.services.legal_data import docx_text, normalize_order_data, validate_amounts


def _pdf(path: Path, text: str) -> None:
    fitz = pytest.importorskip("fitz")
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), text)
    document.save(path)
    document.close()


def _base_data(**updates) -> dict:
    data = {
        "court_name": "Судебный участок № 61 Тверской области",
        "court_address": "170000, г. Тверь, ул. Судебная, д. 1",
        "judge_name": "Мельницкий Д. В.",
        "debtor_full_name": "Сизых Инна Юрьевна",
        "debtor_address": "зарегистрированному в городе Ессентуки, ул. Володарского д. 14, кв. 9",
        "creditor_name": "НАО ПКО «Первое клиентское бюро»",
        "creditor_address": "107061, в городе Москва, Преображенская пл., дом № 8, реквизиты: ИНН 123",
        "case_number": "2-171/2026",
        "uid": "69MS0001-01-2026-000171-01",
        "order_date": "2026-07-03",
        "debt_contract": "кредитный договор № 2026-12390044 от 04.02.2026",
        "debt_basis_type": "credit_agreement",
        "debt_basis_number": "2026-12390044",
        "debt_basis_date": "2026-02-04",
        "debt_period": "с 2026-02-04 по 2026-03-05",
        "debt_amount": "10 000 руб. 00 коп.",
        "state_duty": "500 руб. 00 коп.",
        "total_amount": "10 500 руб. 00 коп.",
    }
    data.update(updates)
    normalized = normalize_order_data(data)
    # Structured fields are authoritative over fallback debt_contract parsing.
    for key in ("debt_basis_type", "debt_basis_number", "debt_basis_date"):
        if key in data:
            normalized[key] = data[key]
    return normalized


def _artifacts(tmp_path: Path, data: dict, *, restore_reason: str | None = None):
    received = date(2026, 7, 10)
    deadline = date(2099, 7, 20)
    ctx = StatementContext(
        data=data,
        received_date=received,
        deadline_date=deadline,
        document_date=date(2026, 7, 11),
        restore_reason=restore_reason,
    )
    full_docx = tmp_path / "full.docx"
    instruction = tmp_path / "instruction.docx"
    full_pdf = tmp_path / "full.pdf"
    preview = tmp_path / "preview.pdf"
    _render_statement_docx(full_docx, ctx, StyleProfile.normal())
    _build_instruction_doc(instruction, deadline="20 июля 2099 года", restore_term=False)
    _pdf(full_pdf, "clean full output")
    _pdf(preview, "redacted preview")
    return ctx, full_docx, full_pdf, preview, instruction


def _qa(tmp_path: Path, data: dict, **kwargs):
    ctx, full_docx, full_pdf, preview, instruction = _artifacts(tmp_path, data)
    return run_document_qa(
        data=data,
        received_date=ctx.received_date,
        deadline_date=ctx.deadline_date,
        full_docx=kwargs.get("full_docx", full_docx),
        full_pdf=kwargs.get("full_pdf", full_pdf),
        preview_pdf=kwargs.get("preview_pdf", preview),
        instruction_docx=kwargs.get("instruction_docx", instruction),
        require_preview_pdf=kwargs.get("require_preview_pdf", True),
        restore_reason=kwargs.get("restore_reason"),
        amount_check=validate_amounts(data),
    )


def test_case_155_structured_name_passes_without_case_heuristic(tmp_path):
    qa = _qa(tmp_path, _base_data(debtor_full_name="Сизых Инна Юрьевна"))
    assert qa.ok, qa.blocking_errors


def test_case_169_dates_contract_number_and_integrity_pass(tmp_path):
    data = _base_data()
    qa = _qa(tmp_path, data)
    text = docx_text(str(tmp_path / "full.docx"))
    assert "2026-12390044" in text
    assert "4 февраля 2026 года" in text
    assert "2026-02-04" not in text
    assert "по по" not in text
    assert qa.ok, qa.blocking_errors


def test_case_171_equivalent_addresses_pass_but_other_house_fails(tmp_path):
    data = _base_data()
    qa = _qa(tmp_path, data)
    assert qa.ok, qa.blocking_errors
    rendered = docx_text(str(tmp_path / "full.docx"))
    assert canonicalize_address_for_qa("г. Ессентуки, ул. Володарского, д. 14, кв. 9") in canonicalize_address_for_qa(rendered)
    assert canonicalize_address_for_qa("107061, г. Москва, Преображенская пл., д. 8") in canonicalize_address_for_qa(rendered)

    wrong = _base_data(debtor_address="г. Ессентуки, ул. Володарского, д. 15, кв. 9")
    wrong_qa = run_document_qa(
        data=wrong,
        received_date=date(2026, 7, 10),
        deadline_date=date(2099, 7, 20),
        full_docx=tmp_path / "full.docx",
        full_pdf=tmp_path / "full.pdf",
        preview_pdf=tmp_path / "preview.pdf",
        instruction_docx=tmp_path / "instruction.docx",
        require_preview_pdf=True,
        amount_check=validate_amounts(wrong),
    )
    assert any("field=debtor_address" in error for error in wrong_qa.integrity_errors)


def test_optional_preview_may_be_absent(tmp_path):
    data = _base_data()
    ctx, full_docx, full_pdf, _, instruction = _artifacts(tmp_path, data)
    qa = run_document_qa(
        data=data, received_date=ctx.received_date, deadline_date=ctx.deadline_date,
        full_docx=full_docx, full_pdf=full_pdf, preview_pdf=None,
        instruction_docx=instruction, require_preview_pdf=False,
        amount_check=validate_amounts(data),
    )
    assert qa.ok, qa.blocking_errors


@pytest.mark.parametrize("artifact", ["full_docx", "full_pdf"])
def test_corrupt_required_artifact_blocks(tmp_path, artifact):
    data = _base_data()
    ctx, full_docx, full_pdf, preview, instruction = _artifacts(tmp_path, data)
    target = full_docx if artifact == "full_docx" else full_pdf
    target.write_bytes(b"not a valid artifact")
    qa = run_document_qa(
        data=data, received_date=ctx.received_date, deadline_date=ctx.deadline_date,
        full_docx=full_docx, full_pdf=full_pdf, preview_pdf=preview,
        instruction_docx=instruction, amount_check=validate_amounts(data),
    )
    assert any(f"artifact={artifact}" in error for error in qa.artifact_errors)


def test_empty_instruction_blocks(tmp_path):
    data = _base_data()
    ctx, full_docx, full_pdf, preview, instruction = _artifacts(tmp_path, data)
    Document().save(instruction)
    qa = run_document_qa(
        data=data, received_date=ctx.received_date, deadline_date=ctx.deadline_date,
        full_docx=full_docx, full_pdf=full_pdf, preview_pdf=preview,
        instruction_docx=instruction, amount_check=validate_amounts(data),
    )
    assert any("artifact=instruction_docx error=empty_text" in error for error in qa.artifact_errors)


def test_restore_reason_is_blocking_when_deadline_missed(tmp_path):
    data = _base_data()
    ctx, full_docx, full_pdf, preview, instruction = _artifacts(tmp_path, data)
    qa = run_document_qa(
        data=data, received_date=ctx.received_date, deadline_date=date(2020, 1, 1),
        full_docx=full_docx, full_pdf=full_pdf, preview_pdf=preview,
        instruction_docx=instruction, amount_check=validate_amounts(data),
    )
    assert not qa.ok
    assert any("field=restore_reason" in error for error in qa.integrity_errors)


def test_address_selection_is_shared_and_explicit():
    field, value = select_creditor_address_for_render({
        "creditor_legal_address": "г. Москва, ул. Ленина, д. 1",
        "creditor_correspondence_address": "г. Самара, ул. Мира, д. 2",
        "creditor_address": "legacy",
    })
    assert (field, value) == ("creditor_legal_address", "г. Москва, ул. Ленина, д. 1")
    assert select_creditor_address_for_render({"creditor_address": "г. Омск, ул. Лесная, д. 3"})[0] == "creditor_address"


def test_total_amount_requires_explicit_render_mode():
    default_text = " ".join(build_statement_paragraphs(StatementContext(
        data=_base_data(), received_date=date(2026, 7, 10),
        deadline_date=date(2099, 7, 20), document_date=date(2026, 7, 11),
    )))
    explicit_text = " ".join(build_statement_paragraphs(StatementContext(
        data=_base_data(amount_render_mode="explicit_total"),
        received_date=date(2026, 7, 10), deadline_date=date(2099, 7, 20),
        document_date=date(2026, 7, 11),
    )))
    assert "Общая сумма взыскания" not in default_text
    assert "Общая сумма взыскания составляет 10 500 руб. 00 коп." in explicit_text


def test_output_date_rules_do_not_confuse_contract_numbers():
    assert "numeric_date" in output_format_violations("Дата приказа 03.07.2026")
    assert "iso_date" in output_format_violations("Дата приказа 2026-07-03")
    assert "iso_date" not in output_format_violations("Договор № 2026-12390044")


def test_entity_canonicalization_handles_initials_quotes_and_yo():
    assert canonicalize_entity_for_qa("Озёров Р.П. «Банк»") == canonicalize_entity_for_qa('Озеров Р. П. "Банк"')


def test_source_only_passport_word_does_not_block_clean_output(tmp_path):
    data = _base_data(debtor_name_context="паспорт указан в исходном приказе")
    qa = _qa(tmp_path, data)
    assert qa.ok, qa.blocking_errors


def test_real_placeholder_in_full_document_blocks(tmp_path):
    data = _base_data()
    ctx, full_docx, full_pdf, preview, instruction = _artifacts(tmp_path, data)
    document = Document(full_docx)
    document.add_paragraph("MISSING")
    document.save(full_docx)
    qa = run_document_qa(
        data=data, received_date=ctx.received_date, deadline_date=ctx.deadline_date,
        full_docx=full_docx, full_pdf=full_pdf, preview_pdf=preview,
        instruction_docx=instruction, amount_check=validate_amounts(data),
    )
    assert "MISSING" in qa.bad_tokens


def test_debt_contract_fallback_has_no_repeated_prepositions():
    data = _base_data(
        debt_basis_number="",
        debt_basis_date="",
        debt_basis_type="",
        debt_contract="по договору, заключенный с ООО «Банк» от 04.02.2026",
    )
    text = " ".join(build_statement_paragraphs(StatementContext(
        data=data, received_date=date(2026, 7, 10),
        deadline_date=date(2099, 7, 20), document_date=date(2026, 7, 11),
    )))
    assert "по по" not in text
    assert "заключённому с ООО «Банк»" in text


def test_render_contract_covers_all_extracted_fields():
    from app.services.llm import ORDER_SCHEMA_HINT
    from app.services.tesseract_ai import ORDER_FIELD_KEYS, SIMPLE_FIELD_KEYS

    extracted = set(ORDER_SCHEMA_HINT) | set(ORDER_FIELD_KEYS) | set(SIMPLE_FIELD_KEYS)
    covered = {field for spec in RENDER_CONTRACT for field in spec.source_fields}
    uncovered = extracted - covered - set(INTENTIONALLY_NOT_RENDERED_FIELDS)
    assert not uncovered, f"Extracted fields missing render strategy: {sorted(uncovered)}"
