from __future__ import annotations

import re
import textwrap
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from app.config import Settings
from app.models import Case, User
from app.services.document_qa import run_document_qa
from app.services.legal_data import (
    FIELD_LABELS,
    format_money_rub_kop,
    clean_case_number,
    clean_uid,
    is_deadline_missed,
    missing_order_fields,
    normalize_debtor_name_fields,
    normalize_order_data,
    suggest_nominative_full_name,
    validate_before_generation,
)
from app.services.name_normalizer import make_short_name
from app.services.pdf_tools import check_pdf_dependencies, convert_docx_to_pdf, create_preview_pdf, pdf_text
from app.utils import ensure_dir, h, safe_json_loads


DOCUMENT_DIR = Path("storage/documents")


def _required(data: dict, key: str) -> str:
    value = str(data.get(key) or "").strip()
    if not value:
        raise ValueError(f"Missing required document field: {key}")
    return value


def _optional(data: dict, key: str) -> str:
    return str(data.get(key) or "").strip()


def _short_name(full_name: str) -> str:
    return make_short_name(full_name)


def _date_long_text(raw: str) -> str:
    parsed = None
    try:
        from app.utils import parse_russian_date

        parsed = parse_russian_date(raw)
    except Exception:
        parsed = None
    if not parsed:
        return raw
    months = [
        "января",
        "февраля",
        "марта",
        "апреля",
        "мая",
        "июня",
        "июля",
        "августа",
        "сентября",
        "октября",
        "ноября",
        "декабря",
    ]
    return f"{parsed.day} {months[parsed.month - 1]} {parsed.year} года"


def _normalize_creditor_address(address: str) -> str:
    text = re.sub(r"\s+", " ", address).strip(" ,.;")
    text = text.replace("в городе ", "г. ")
    text = re.sub(r"^город\s+", "г. ", text, flags=re.IGNORECASE)
    text = re.sub(r"^г\.\s*Москва", "г. Москва", text, flags=re.IGNORECASE)
    text = re.sub(r"^107061,\s*в\s+г\.?\s*Москва", "107061, г. Москва", text, flags=re.IGNORECASE)
    text = re.sub(r"^107061,\s*в\s+городе\s+Москва", "107061, г. Москва", text, flags=re.IGNORECASE)
    return text


def _normalize_court_for_addressee(court: str) -> str:
    court = court.strip().rstrip(".")
    lower = court.lower()
    if lower.startswith("мировому судье"):
        return court
    if lower.startswith("мировой судья"):
        return "Мировому судье " + court[len("мировой судья") :].strip()
    if lower.startswith("судебный участок"):
        normalized = re.sub(r"^судебный участок", "судебного участка", court, flags=re.IGNORECASE)
        normalized = re.sub(r"№\s*(\d+)", r"№ \1", normalized)
        return "Мировому судье " + normalized
    return court


def _normalize_court_for_body(court: str) -> str:
    court = court.strip().rstrip(".")
    lower = court.lower()
    if lower.startswith("мировому судье"):
        rest = court[len("мировому судье") :].strip()
        return f"мировым судьей {rest}" if rest else "мировым судьей"
    if lower.startswith("мировой судья"):
        rest = court[len("мировой судья") :].strip()
        return f"мировым судьей {rest}" if rest else "мировым судьей"
    if lower.startswith("судебный участок"):
        normalized = re.sub(r"^судебный участок", "судебного участка", court, flags=re.IGNORECASE)
        normalized = re.sub(r"№\s*(\d+)", r"№ \1", normalized)
        return f"мировым судьей {normalized}"
    return court


def _court_instrumental(court: str) -> str:
    return _normalize_court_for_body(court)


def _case_identifier(data: dict) -> str:
    case_number = clean_case_number(_required(data, "case_number"))
    uid = clean_uid(_optional(data, "uid"))
    parts = [f"№ {case_number}"]
    if uid:
        parts.append(f"УИД {uid}")
    return ", ".join(parts)


def _case_identifier_short(data: dict) -> str:
    case_number = clean_case_number(_required(data, "case_number"))
    uid = clean_uid(_optional(data, "uid"))
    if uid:
        return f"№ {case_number}, УИД {uid}"
    return f"№ {case_number}"


def _contract_phrase(value: str) -> str:
    value = re.sub(r"\s+", " ", value).strip(" ,.;")
    lower = value.lower()
    if lower.startswith("по "):
        return value
    if lower.startswith(("договор", "кредит", "карта", "счет", "счёт")):
        return f"по {value}"
    if value.startswith("№"):
        return f"по договору {value}"
    return f"по договору {value}"


def _period_phrase(value: str) -> str:
    value = re.sub(r"\s+", " ", value).strip(" ,.;")
    value = re.sub(r"^(за\s+период|период|за)\s+", "", value, flags=re.IGNORECASE).strip(" ,.;")
    return f"за период {value}"


def _money_sentence(data: dict) -> str:
    debt = _required(data, "debt_amount")
    state_duty = _optional(data, "state_duty")
    total = _optional(data, "total_amount")
    parts = [f"задолженности в размере {debt}"]
    if state_duty:
        parts.append(f"расходов по оплате государственной пошлины в размере {state_duty}")
    if total:
        parts.append(f"всего {total}")
    return ", ".join(parts)


def _set_cell_borderless(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def _setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15


def _paragraph(
    doc: Document,
    text: str = "",
    *,
    bold: bool = False,
    size: int | None = None,
    align=None,
    before=0,
    after=6,
    first_line=False,
):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.15
    if first_line:
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    if size:
        run.font.size = Pt(size)
    return p


def _address_block(doc: Document, data: dict, user: User) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    table.columns[0].width = Cm(7.0)
    table.columns[1].width = Cm(9.5)
    left, right = table.rows[0].cells
    _set_cell_borderless(left)
    _set_cell_borderless(right)
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    debtor_full_name = _required(data, "debtor_full_name")
    court_addressee = data.get("court_addressee") or _normalize_court_for_addressee(_required(data, "court_name"))
    lines = [
        court_addressee,
        _required(data, "court_address"),
        "",
        "Должник:",
        debtor_full_name,
        f"адрес: {_required(data, 'debtor_address')}",
    ]
    lines.extend(
        [
            "",
            f"Взыскатель: {_required(data, 'creditor_name')}",
            _normalize_creditor_address(_required(data, "creditor_address")),
        ]
    )
    case_number = clean_case_number(_required(data, "case_number"))
    uid = clean_uid(_optional(data, "uid"))
    lines.extend(["", f"Дело/производство № {case_number}"])
    if uid:
        lines.append(f"УИД: {uid}")
    for line in lines:
        p = right.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if line:
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.name = "Times New Roman"
    doc.add_paragraph()


def _meta_line(doc: Document, data: dict) -> None:
    p = _paragraph(doc, f"Дело/производство {_case_identifier(data)}", size=11, align=WD_ALIGN_PARAGRAPH.RIGHT, after=10)
    for run in p.runs:
        run.font.color.rgb = RGBColor(90, 90, 90)


def build_statement_paragraphs(data: dict, received_date: date, deadline_date: date | None, restore_reason: str | None = None) -> list[str]:
    court_body = data.get("court_instrumental") or _court_instrumental(_required(data, "court_name"))
    case_identifier = _case_identifier_short(data)
    order_date = _required(data, "order_date")
    order_date_long = _date_long_text(order_date)
    creditor = _required(data, "creditor_name")
    contract = _contract_phrase(_required(data, "debt_contract"))
    period = _period_phrase(_required(data, "debt_period"))
    received_long = _date_long_text(received_date.strftime("%d.%m.%Y"))
    deadline = deadline_date.strftime("%d.%m.%Y") if deadline_date else ""
    money_part = _money_sentence(data)
    base = [
        f"{order_date_long} {court_body} вынесен судебный приказ по делу/производству {case_identifier}, о взыскании с меня в пользу {creditor} {money_part} {contract} {period}.",
        f"Копия судебного приказа получена мной {received_long}.",
    ]
    if is_deadline_missed(deadline_date):
        reason = restore_reason or "Причина пропуска срока не указана."
        return base + [
            f"Десятидневный срок подачи возражений истек {deadline}. {reason}",
            "Прошу восстановить пропущенный процессуальный срок.",
            "С судебным приказом не согласен, возражаю относительно его исполнения в полном объеме.",
            "На основании изложенного, руководствуясь статьями 112, 128, 129 ГПК РФ,",
            "ПРОШУ:",
            f"1. Восстановить срок для подачи возражений относительно исполнения судебного приказа от {order_date_long}, вынесенного {court_body} по делу/производству {case_identifier}.",
            f"2. Отменить судебный приказ от {order_date_long}, вынесенный {court_body} по делу/производству {case_identifier}.",
            "3. Направить мне копию определения об отмене судебного приказа по адресу, указанному в настоящих возражениях.",
        ]
    return base + [
        "С судебным приказом не согласен, возражаю относительно его исполнения в полном объеме.",
        "Настоящие возражения подаются в установленный законом срок.",
        "На основании изложенного, руководствуясь статьями 128, 129 ГПК РФ,",
        "ПРОШУ:",
        f"1. Отменить судебный приказ от {order_date_long}, вынесенный {court_body} по делу/производству {case_identifier}.",
        "2. Направить мне копию определения об отмене судебного приказа по адресу, указанному в настоящих возражениях.",
    ]


def _redacted_line(doc: Document, prefix: str = "") -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if prefix:
        p.add_run(prefix + " ")
    run = p.add_run("▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒ ▒▒▒▒▒▒▒▒▒▒▒▒▒▒ ▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒▒")
    run.font.name = "Times New Roman"
    run.font.color.rgb = RGBColor(185, 185, 185)
    run.font.size = Pt(12)


def _add_preview_text(doc: Document, text: str, line_no: int) -> int:
    chunks = textwrap.wrap(text, width=92, break_long_words=False, break_on_hyphens=False) or [text]
    for chunk in chunks:
        match = re.match(r"^(\d+\.)\s+", chunk)
        if line_no % 2 == 0:
            _redacted_line(doc, match.group(1) if match else "")
        else:
            _paragraph(doc, chunk, first_line=False, after=2)
        line_no += 1
    return line_no


def _add_body(doc: Document, paragraphs: list[str], *, preview: bool, restore_term: bool) -> None:
    _paragraph(doc, "ВОЗРАЖЕНИЯ", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, before=4, after=2)
    subtitle = "относительно исполнения судебного приказа с ходатайством о восстановлении срока" if restore_term else "относительно исполнения судебного приказа"
    _paragraph(doc, subtitle, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=10)
    _paragraph(doc, "/ заявление об отмене судебного приказа /", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, after=12)
    line_no = 1
    for text in paragraphs:
        if text == "ПРОШУ:":
            _paragraph(doc, text, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, before=6, after=6)
            continue
        if preview:
            line_no = _add_preview_text(doc, text, line_no)
        else:
            _paragraph(doc, text, first_line=True, after=7)


def _add_attachments_and_signature(doc: Document, data: dict, debtor: str, document_date: date, *, preview: bool, restore_term: bool) -> None:
    _paragraph(doc, "Приложения:", bold=True, before=6, after=4)
    items = [
        f"Копия судебного приказа от {_required(data, 'order_date')}.",
        "Копия конверта или иной документ, подтверждающий дату получения судебного приказа.",
        "Копия настоящего заявления для взыскателя.",
    ]
    if restore_term:
        items.insert(2, "Документы, подтверждающие дату фактического получения судебного приказа и причины пропуска срока.")
    line_no = 1
    for i, text in enumerate(items, 1):
        if preview:
            line_no = _add_preview_text(doc, f"{i}. {text}", line_no)
        else:
            _paragraph(doc, f"{i}. {text}", after=3)
    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(8.5)
    table.columns[1].width = Cm(7.0)
    for cell in table.rows[0].cells:
        _set_cell_borderless(cell)
    table.rows[0].cells[0].paragraphs[0].add_run(f"«{document_date.day:02d}» {['января','февраля','марта','апреля','мая','июня','июля','августа','сентября','октября','ноября','декабря'][document_date.month - 1]} {document_date.year} г.")
    right = table.rows[0].cells[1].paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    right.add_run(f"_____________ /{_short_name(debtor)}/")
    if not preview:
        pass
    if preview:
        p = _paragraph(doc, "Предпросмотр: каждая вторая строка скрыта до оплаты.", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=0)
        p.runs[0].font.color.rgb = RGBColor(120, 120, 120)


def _build_instruction_doc(path: Path, *, deadline: str, restore_term: bool) -> None:
    doc = Document()
    _setup_styles(doc)
    _paragraph(doc, "Инструкция", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    _paragraph(doc, "по подаче возражений в суд", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    lines = [
        f"Срок на подачу: до {deadline} включительно.",
        "Документ можно подать лично в канцелярию мирового судьи.",
        "Документ можно отправить заказным письмом с описью вложения.",
        "Если отправка почтой, важно сдать письмо до 24:00 последнего дня срока.",
        "Сохраните чек, опись и трек-номер.",
        "Подпись в заявлении ставится от руки синей ручкой.",
    ]
    if restore_term:
        lines.insert(1, "Так как срок пропущен, в заявлении уже включено ходатайство о восстановлении срока.")
    for line in lines:
        _paragraph(doc, line, first_line=True, after=6)
    doc.save(path)


def _validate_package_texts(
    full_docx: Path,
    full_pdf: Path | None,
    *,
    preview_pdf: Path | None = None,
    preview_docx: Path | None = None,
    card_text: str,
    data: dict,
    received_date: date | None,
    deadline_date: date | None,
    instruction_path: Path | None,
    restore_reason: str | None,
    require_preview_pdf: bool,
) -> None:
    qa = run_document_qa(
        data=data,
        received_date=received_date,
        deadline_date=deadline_date,
        full_docx=full_docx,
        full_pdf=full_pdf,
        preview_pdf=preview_pdf,
        instruction_docx=instruction_path,
        preview_docx=preview_docx,
        card_text=card_text,
        restore_reason=restore_reason,
        require_preview_pdf=require_preview_pdf,
    )
    if not qa.ok:
        raise ValueError("Документ не прошел QA: " + "; ".join(qa.reasons or qa.bad_tokens))


def create_case_documents(case: Case, user: User, settings: Settings, *, restore_reason: str | None = None) -> tuple[Path, Path | None, Path | None, Path | None, Path]:
    ensure_dir(DOCUMENT_DIR)
    case_dir = ensure_dir(DOCUMENT_DIR / f"case_{case.id}")
    data = normalize_order_data(safe_json_loads(case.extracted_json, {}))
    validation = validate_before_generation(data, case.received_date)
    if not validation.ok:
        labels = [FIELD_LABELS.get(field, field) for field in validation.missing]
        raise ValueError("Нельзя сформировать заявление: " + ", ".join(labels))
    if not case.received_date:
        raise ValueError("Нельзя сформировать заявление без даты получения")
    deps_ok, dep_errors = check_pdf_dependencies(require_preview_pdf_for_payment=settings.require_pdf_preview_for_payment)
    if not deps_ok and settings.require_pdf_preview_for_payment:
        raise ValueError("; ".join(dep_errors))

    restore_term = is_deadline_missed(case.deadline_date)
    debtor = _required(data, "debtor_full_name")
    paragraphs = build_statement_paragraphs(data, case.received_date, case.deadline_date, restore_reason=restore_reason)
    suffix = "restore_term" if restore_term else "in_time"
    full_docx = case_dir / f"statement_{suffix}_{case.id}.docx"
    full_pdf = case_dir / f"statement_{suffix}_{case.id}.pdf"
    preview_pdf = case_dir / f"preview_statement_{suffix}_{case.id}.pdf"
    preview_docx = case_dir / f"preview_statement_{suffix}_{case.id}.docx"
    instruction_path = case_dir / f"instruction_{case.id}.docx"
    document_date = date.today()

    doc = Document()
    _setup_styles(doc)
    _address_block(doc, data, user)
    _add_body(doc, paragraphs, preview=False, restore_term=restore_term)
    _add_attachments_and_signature(doc, data, debtor, document_date, preview=False, restore_term=restore_term)
    doc.save(full_docx)
    pdf_conversion_failed = False
    try:
        full_pdf = convert_docx_to_pdf(
            full_docx,
            case_dir,
            allow_dev_fallback=settings.allow_dev_docx_preview,
        )
    except Exception:
        full_pdf = None
        pdf_conversion_failed = True

    preview_pdf_path: Path | None = None
    preview_docx_path: Path | None = None
    use_pdf_preview = settings.enable_pdf_preview and settings.document_preview_mode != "docx"
    if use_pdf_preview and full_pdf is not None:
        preview_pdf_path = create_preview_pdf(full_pdf, preview_pdf)
    elif settings.allow_dev_docx_preview and (settings.document_preview_mode == "docx" or pdf_conversion_failed):
        preview_doc = Document()
        _setup_styles(preview_doc)
        _address_block(preview_doc, data, user)
        _add_body(preview_doc, paragraphs, preview=True, restore_term=restore_term)
        _add_attachments_and_signature(preview_doc, data, debtor, document_date, preview=True, restore_term=restore_term)
        preview_doc.save(preview_docx)
        preview_docx_path = preview_docx

    deadline = case.deadline_date.strftime("%d.%m.%Y") if case.deadline_date else "уточняется"
    _build_instruction_doc(instruction_path, deadline=deadline, restore_term=restore_term)

    card_text = extraction_preview(data, case.received_date, [], case.deadline_date, include_name_debug=False)
    require_preview_pdf = settings.require_pdf_preview_for_payment and use_pdf_preview
    _validate_package_texts(
        full_docx,
        full_pdf,
        preview_pdf=preview_pdf_path,
        preview_docx=preview_docx_path,
        card_text=card_text,
        data=data,
        received_date=case.received_date,
        deadline_date=case.deadline_date,
        instruction_path=instruction_path,
        restore_reason=restore_reason,
        require_preview_pdf=require_preview_pdf,
    )

    return full_docx, full_pdf, preview_pdf_path, preview_docx_path, instruction_path


def extraction_preview(
    data: dict,
    received_date: date | None,
    missing: list[str],
    deadline_date: date | None = None,
    *,
    include_name_debug: bool = True,
) -> str:
    data = normalize_order_data(data)
    lines = [
        "🔎 <b>Проверьте данные</b>",
        "",
        f"<b>Суд:</b> {h(data.get('court_name') or 'не заполнено')}",
        f"<b>Адрес суда:</b> {h(data.get('court_address') or 'не заполнено')}",
        f"<b>Должник:</b> {h(data.get('debtor_full_name') or 'не заполнено')}",
        f"<b>Адрес должника:</b> {h(data.get('debtor_address') or 'не заполнено')}",
        f"<b>Взыскатель:</b> {h(data.get('creditor_name') or 'не заполнено')}",
        f"<b>Номер дела:</b> {h(data.get('case_number') or 'не заполнено')}",
        f"<b>УИД:</b> {h(data.get('uid') or 'нет в приказе')}",
        f"<b>Дата приказа:</b> {h(data.get('order_date') or 'не заполнено')}",
        f"<b>Договор:</b> {h(data.get('debt_contract') or 'не заполнено')}",
        f"<b>Период:</b> {h(data.get('debt_period') or 'не заполнено')}",
        f"<b>Сумма долга:</b> {h(data.get('debt_amount') or 'не заполнено')}",
        f"<b>Госпошлина:</b> {h(data.get('state_duty') or 'не указана')}",
        f"<b>Дата получения:</b> {received_date.strftime('%d.%m.%Y') if received_date else 'не указана'}",
    ]
    if deadline_date:
        lines.append(f"<b>Срок до:</b> {deadline_date.strftime('%d.%m.%Y')} включительно")
    if include_name_debug:
        raw_name = data.get("debtor_name_raw") or ""
        if raw_name and raw_name != data.get("debtor_full_name"):
            lines.append(f"<i>Исходно распознано:</i> {h(raw_name)}")
            lines.append(f"<i>Нормализовано:</i> {h(data.get('debtor_full_name') or '')}")
    if missing:
        labels = [FIELD_LABELS.get(field, field) for field in missing]
        lines.extend(["", "⚠️ <b>Перед генерацией нужно заполнить:</b>", ", ".join(labels)])
    else:
        lines.extend(["", "Если все верно, можно готовить документы. Если видите ошибку OCR, исправьте поле кнопкой ниже."])
    return "\n".join(lines)
