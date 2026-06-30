from __future__ import annotations

from dataclasses import dataclass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

A4_WIDTH = Cm(21)
A4_HEIGHT = Cm(29.7)
MARGIN_LEFT = Cm(2.5)
MARGIN_RIGHT = Cm(1.5)
MARGIN_TOP = Cm(2.0)
MARGIN_BOTTOM = Cm(2.0)
COMPACT_MARGIN_BOTTOM = Cm(1.5)

FONT_NAME = "Times New Roman"


@dataclass(frozen=True)
class StyleProfile:
    body_font_size: float = 13.0
    header_font_size: float = 11.0
    title_font_size: float = 14.0
    subtitle_font_size: float = 12.0
    section_font_size: float = 13.0
    line_spacing: float = 1.08
    body_space_after: float = 4.0
    header_space_after: float = 0.0
    header_block_space: float = 5.0
    title_space_before: float = 8.0
    title_space_after: float = 2.0
    subtitle_space_after: float = 6.0
    title_note_space_after: float = 6.0
    list_space_after: float = 3.0
    compact: bool = False

    @classmethod
    def normal(cls) -> StyleProfile:
        return cls()

    @classmethod
    def compact(cls) -> StyleProfile:
        return cls(
            body_font_size=12.5,
            header_font_size=10.5,
            line_spacing=1.0,
            body_space_after=2.0,
            header_block_space=2.0,
            title_space_before=4.0,
            title_space_after=1.0,
            subtitle_space_after=3.0,
            title_note_space_after=3.0,
            list_space_after=2.0,
            compact=True,
        )

    @classmethod
    def ultra_compact(cls) -> StyleProfile:
        return cls(
            body_font_size=12.0,
            header_font_size=10.0,
            title_font_size=13.5,
            subtitle_font_size=11.5,
            section_font_size=12.0,
            line_spacing=0.96,
            body_space_after=1.0,
            header_block_space=1.0,
            title_space_before=2.0,
            title_space_after=0.0,
            subtitle_space_after=2.0,
            title_note_space_after=2.0,
            list_space_after=1.0,
            compact=True,
        )


def setup_page(doc: Document, profile: StyleProfile | None = None) -> None:
    section = doc.sections[0]
    section.page_width = A4_WIDTH
    section.page_height = A4_HEIGHT
    section.top_margin = MARGIN_TOP
    section.bottom_margin = COMPACT_MARGIN_BOTTOM if profile and profile.compact else MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT


def _set_run_font(run, *, size: float, bold: bool = False) -> None:
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold


def add_paragraph(
    doc: Document,
    text: str,
    profile: StyleProfile,
    *,
    bold: bool = False,
    size: float | None = None,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    before: float = 0,
    after: float | None = None,
    first_line_indent: float | None = None,
    left_indent: float | None = None,
    hanging_indent: float | None = None,
    keep_with_next: bool = False,
    keep_together: bool = False,
    line_spacing: float | None = None,
):
    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after if after is not None else profile.body_space_after)
    fmt.line_spacing = line_spacing if line_spacing is not None else profile.line_spacing
    fmt.keep_with_next = keep_with_next
    fmt.keep_together = keep_together
    p.alignment = align
    if first_line_indent is not None:
        fmt.first_line_indent = Cm(first_line_indent)
    if left_indent is not None:
        fmt.left_indent = Cm(left_indent)
    if hanging_indent is not None:
        fmt.first_line_indent = Cm(-hanging_indent)
    run = p.add_run(text)
    _set_run_font(run, size=size or profile.body_font_size, bold=bold)
    return p


def page_margins_cm(profile: StyleProfile | None = None) -> dict[str, float]:
    return {
        "left": 2.5,
        "right": 1.5,
        "top": 2.0,
        "bottom": 1.5 if profile and profile.compact else 2.0,
    }
