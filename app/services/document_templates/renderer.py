from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from app.config import Settings
from app.models import Case, User
from app.services.document_qa import run_document_qa
from app.services.document_templates.statement_templates import (
    StatementContext,
    build_attachments,
    build_header_lines,
    build_statement_paragraphs,
    date_long_text,
    debtor_short_name,
    signature_date_text,
)
from app.services.document_templates.styles import (
    A4_HEIGHT,
    A4_WIDTH,
    FONT_NAME,
    StyleProfile,
    add_paragraph,
    page_margins_cm,
    setup_page,
)
from app.services.document_visual_qa import VisualQAResult, run_visual_qa
from app.services.legal_data import (
    FIELD_LABELS,
    is_deadline_missed,
    missing_order_fields,
    normalize_order_data,
    validate_amounts,
    validate_before_generation,
)
from app.services.pdf_tools import check_pdf_dependencies, convert_docx_to_pdf, create_preview_pdf, pdf_page_count
from app.utils import ensure_dir, safe_json_loads


DOCUMENT_DIR = Path("storage/documents")


@dataclass
class DocumentArtifacts:
    full_docx_path: Path
    full_pdf_path: Path | None
    preview_pdf_path: Path | None
    instruction_docx_path: Path
    qa_report: dict
    visual_qa: VisualQAResult | None = None


def _set_cell_borderless(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        element.set(qn("w:val"), "nil")


def _add_header_block(doc: Document, ctx: StatementContext, profile: StyleProfile) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    table.autofit = False
    table.columns[0].width = Cm(7.0)
    table.columns[1].width = Cm(9.8)
    left, right = table.rows[0].cells
    _set_cell_borderless(left)
    _set_cell_borderless(right)
    right.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    prev_blank = False
    for line in build_header_lines(ctx):
        if not line:
            if prev_blank:
                continue
            p = right.add_paragraph()
            p.paragraph_format.space_after = Pt(profile.header_block_space)
            prev_blank = True
            continue
        prev_blank = False
        p = right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(profile.header_space_after)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
        run.font.size = Pt(profile.header_font_size)
    doc.add_paragraph()


def _add_title_block(doc: Document, ctx: StatementContext, profile: StyleProfile) -> None:
    restore_term = is_deadline_missed(ctx.deadline_date, ctx.document_date)
    add_paragraph(
        doc,
        "ВОЗРАЖЕНИЯ",
        profile,
        bold=True,
        size=profile.title_font_size,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        before=profile.title_space_before,
        after=profile.title_space_after,
    )
    subtitle = (
        "относительно исполнения судебного приказа\nс ходатайством о восстановлении срока"
        if restore_term
        else "относительно исполнения судебного приказа"
    )
    for index, line in enumerate(subtitle.split("\n")):
        add_paragraph(
            doc,
            line,
            profile,
            size=profile.subtitle_font_size,
            align=WD_ALIGN_PARAGRAPH.CENTER,
            after=2 if index == 0 and restore_term else profile.subtitle_space_after,
        )
    add_paragraph(
        doc,
        "(заявление об отмене судебного приказа)",
        profile,
        size=11,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        after=profile.title_note_space_after,
    )


def _add_body(doc: Document, paragraphs: list[str], profile: StyleProfile) -> None:
    for text in paragraphs:
        if text == "ПРОШУ:":
            add_paragraph(
                doc,
                text,
                profile,
                bold=True,
                size=profile.section_font_size,
                before=6,
                after=4,
                keep_with_next=True,
                keep_together=True,
            )
            continue
        if re.match(r"^\d+\.", text):
            add_paragraph(
                doc,
                text,
                profile,
                left_indent=0.6,
                hanging_indent=0.4,
                after=profile.list_space_after,
                keep_together=True,
            )
            continue
        add_paragraph(
            doc,
            text,
            profile,
            first_line_indent=1.25,
            after=profile.body_space_after,
        )


def _add_attachments_and_signature(doc: Document, ctx: StatementContext, profile: StyleProfile) -> None:
    add_paragraph(
        doc,
        "Приложения:",
        profile,
        bold=True,
        size=profile.section_font_size,
        before=6,
        after=4,
        keep_with_next=True,
        keep_together=True,
    )
    for index, text in enumerate(build_attachments(ctx), 1):
        add_paragraph(
            doc,
            f"{index}. {text}",
            profile,
            left_indent=0.6,
            hanging_indent=0.4,
            after=profile.list_space_after,
            keep_together=True,
        )
    spacer = doc.add_paragraph()
    spacer_fmt = spacer.paragraph_format
    spacer_fmt.space_before = Pt(0)
    spacer_fmt.space_after = Pt(0)
    spacer_fmt.line_spacing = 1.0
    spacer_run = spacer.add_run(" ")
    spacer_run.font.name = FONT_NAME
    spacer_run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    spacer_run.font.size = Pt(profile.body_font_size)

    p = doc.add_paragraph()
    fmt = p.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing = 1.0
    fmt.tab_stops.add_tab_stop(Cm(8.3), WD_TAB_ALIGNMENT.CENTER)
    fmt.tab_stops.add_tab_stop(Cm(16.8), WD_TAB_ALIGNMENT.RIGHT)
    run = p.add_run(f"{signature_date_text(ctx.document_date)}	_____________	/{debtor_short_name(ctx.data)}/")
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_NAME)
    run.font.size = Pt(profile.body_font_size)


def _build_instruction_doc(path: Path, *, deadline: str, restore_term: bool) -> None:
    profile = StyleProfile.normal()
    doc = Document()
    setup_page(doc)
    add_paragraph(doc, "Инструкция", profile, bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)
    add_paragraph(doc, "по подаче возражений в суд", profile, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, after=14)
    lines = [
        f"Срок на подачу: до {deadline} включительно.",
        "Документ можно подать лично в канцелярию мирового судьи.",
        "Документ можно отправить заказным письмом с описью вложения.",
        "Если отправка почтой, важно сдать письмо до 24:00 последнего дня срока.",
        "Сохраните чек, опись и трек-номер.",
        "Распечатайте документ и поставьте подпись от руки синей ручкой.",
    ]
    if restore_term:
        lines.insert(1, "Так как срок пропущен, в заявлении уже включено ходатайство о восстановлении срока.")
    for line in lines:
        add_paragraph(doc, line, profile, first_line_indent=1.25, after=6)
    doc.save(path)


def _render_statement_docx(path: Path, ctx: StatementContext, profile: StyleProfile) -> None:
    doc = Document()
    setup_page(doc, profile)
    _add_header_block(doc, ctx, profile)
    _add_title_block(doc, ctx, profile)
    paragraphs = build_statement_paragraphs(ctx)
    _add_body(doc, paragraphs, profile)
    _add_attachments_and_signature(doc, ctx, profile)
    doc.save(path)


def _validate_a4_margins(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    section = doc.sections[0]
    if abs(section.page_width - A4_WIDTH) > Cm(0.05):
        raise ValueError("page_width is not A4")
    if abs(section.page_height - A4_HEIGHT) > Cm(0.05):
        raise ValueError("page_height is not A4")


def create_case_documents(
    case: Case,
    user: User,
    settings: Settings,
    *,
    restore_reason: str | None = None,
) -> DocumentArtifacts:
    ensure_dir(DOCUMENT_DIR)
    case_dir = ensure_dir(DOCUMENT_DIR / f"case_{case.id}")
    data = normalize_order_data(safe_json_loads(case.extracted_json, {}))
    validation = validate_before_generation(data, case.received_date)
    if not validation.ok:
        labels = [FIELD_LABELS.get(field, field) for field in validation.missing]
        raise ValueError("Нельзя сформировать заявление: " + ", ".join(labels))
    if not case.received_date:
        raise ValueError("Нельзя сформировать заявление без даты получения")

    amount_check = validate_amounts(data)
    if not amount_check.ok:
        raise ValueError("Суммы требуют проверки: " + "; ".join(amount_check.errors))

    restore_term = is_deadline_missed(case.deadline_date)
    if restore_term and not restore_reason:
        raise ValueError("Срок пропущен, но не указана причина восстановления")

    deps_ok, dep_errors = check_pdf_dependencies(require_preview_pdf_for_payment=settings.require_pdf_preview_for_payment)
    if not deps_ok and settings.require_pdf_preview_for_payment:
        raise ValueError("; ".join(dep_errors))

    document_date = date.today()
    ctx = StatementContext(
        data=data,
        received_date=case.received_date,
        deadline_date=case.deadline_date,
        document_date=document_date,
        restore_reason=restore_reason,
        has_envelope=bool(case.envelope_photo_path),
        manual_date_only=bool(case.received_date and not case.envelope_photo_path),
    )

    suffix = "restore_term" if restore_term else "in_time"
    full_docx = case_dir / f"statement_{suffix}_{case.id}.docx"
    full_pdf = case_dir / f"statement_{suffix}_{case.id}.pdf"
    preview_pdf = case_dir / f"preview_statement_{suffix}_{case.id}.pdf"
    instruction_path = case_dir / f"instruction_{case.id}.docx"

    profile = StyleProfile.normal()
    _render_statement_docx(full_docx, ctx, profile)
    _validate_a4_margins(full_docx)

    full_pdf_path: Path | None = None
    try:
        full_pdf_path = convert_docx_to_pdf(
            full_docx,
            case_dir,
            allow_dev_fallback=settings.allow_dev_docx_preview,
        )
    except Exception as exc:
        if settings.require_pdf_preview_for_payment:
            raise ValueError(f"Не удалось создать PDF: {exc}") from exc

    page_count = pdf_page_count(full_pdf_path) if full_pdf_path else None
    if not restore_term and page_count and page_count > 1:
        profile = StyleProfile.compact()
        _render_statement_docx(full_docx, ctx, profile)
        _validate_a4_margins(full_docx)
        if full_pdf_path:
            full_pdf_path = convert_docx_to_pdf(
                full_docx,
                case_dir,
                allow_dev_fallback=settings.allow_dev_docx_preview,
            )
            page_count = pdf_page_count(full_pdf_path)
        if page_count and page_count > 1:
            profile = StyleProfile.ultra_compact()
            _render_statement_docx(full_docx, ctx, profile)
            _validate_a4_margins(full_docx)
            if full_pdf_path:
                full_pdf_path = convert_docx_to_pdf(
                    full_docx,
                    case_dir,
                    allow_dev_fallback=settings.allow_dev_docx_preview,
                )
                page_count = pdf_page_count(full_pdf_path)

    preview_pdf_path: Path | None = None
    if settings.enable_pdf_preview and full_pdf_path is not None:
        preview_pdf_path = create_preview_pdf(full_pdf_path, preview_pdf)

    deadline = case.deadline_date.strftime("%d.%m.%Y") if case.deadline_date else "уточняется"
    _build_instruction_doc(instruction_path, deadline=deadline, restore_term=restore_term)

    from app.services.documents import extraction_preview

    card_text = extraction_preview(data, case.received_date, [], case.deadline_date, include_name_debug=False)
    require_preview_pdf = settings.require_pdf_preview_for_payment and settings.enable_pdf_preview
    qa = run_document_qa(
        data=data,
        received_date=case.received_date,
        deadline_date=case.deadline_date,
        full_docx=full_docx,
        full_pdf=full_pdf_path,
        preview_pdf=preview_pdf_path,
        instruction_docx=instruction_path,
        card_text=card_text,
        restore_reason=restore_reason,
        require_preview_pdf=require_preview_pdf,
        amount_check=amount_check,
    )
    visual_qa = run_visual_qa(
        full_docx=full_docx,
        full_pdf=full_pdf_path,
        preview_pdf=preview_pdf_path,
        data=data,
        restore_term=restore_term,
        amount_check=amount_check,
        profile=profile,
    )
    qa_report = {
        "document_qa_ok": qa.ok,
        "document_qa_errors": qa.reasons,
        "document_qa_bad_tokens": qa.bad_tokens,
        "visual_qa_ok": visual_qa.ok,
        "visual_qa_errors": visual_qa.errors,
        "visual_qa_warnings": visual_qa.warnings,
        "page_count": visual_qa.page_count,
        "font_name": visual_qa.font_name,
        "body_font_size": visual_qa.body_font_size,
        "margins": visual_qa.margins,
        "amounts": visual_qa.amounts,
    }
    if not qa.ok:
        raise ValueError("Документ не прошел QA: " + "; ".join(qa.reasons or qa.bad_tokens))
    if not visual_qa.ok:
        raise ValueError("Документ не прошел visual QA: " + "; ".join(visual_qa.errors))

    return DocumentArtifacts(
        full_docx_path=full_docx,
        full_pdf_path=full_pdf_path,
        preview_pdf_path=preview_pdf_path,
        instruction_docx_path=instruction_path,
        qa_report=qa_report,
        visual_qa=visual_qa,
    )
